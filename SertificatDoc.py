"""Cоздание документа doc на основе данных из гугл таблицы"""
import docx
import gspread
import zipfile
import time 
from docx.dml.color import ColorFormat
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from loguru import logger
from oauth2client.service_account import ServiceAccountCredentials

TEST_LIST = [['ШУЕППШ', '10', '10', '8', '8', '10', '10', 'Герасимов Игорь (МР-118), Анатолий Жук (И-113)',]]
TEST_LIST2 = [[  'Кочуков Юрий (ЛТ-118)','Герасимов Игорь (МР-118)', 'Анатолий Жук (И-113)', 'Кочуков Юрий (ЛТ-118)']]
NAME_DOC_OUT = '1' # вторая ячейка в гугл таблице отвечает за название документа
NAME_DOC_INPUT = 'LOWSertificat.docx'

@logger.catch
def auntification_Sheet():
    """Возвращает объект для работы с таблицей гугл"""
    SCOPE = ['https://spreadsheets.google.com/feeds',
            'https://www.googleapis.com/auth/drive'] # что то для чего-то нужно Костыль    
    CREDS = ServiceAccountCredentials.from_json_keyfile_name(
            "/Users/igorgerasimov/Desktop/Мусор/kgtaprojects-af476ce0430b.json", 
            SCOPE) # Секретынй файл json для доступа к API
    # Также нужно добавить в таблицу сервисный аккаунт kgta-34@kgtaprojects.iam.gserviceaccount.com
    CLIENT = gspread.authorize(CREDS)
    return CLIENT

def toFixed(numObj, digits=0):
    return round(numObj, digits)

@logger.catch
def split_List(user: str) -> list:
    users = user.split(',')
    return users

@logger.catch
def add_in_doc(user: list) -> None:
    """Создание док файла для одного человека"""
    allPersent = 0
    furst = 1
    second = 2
 
    for index in range(1,7):
        allPersent += int(user[index])
    if (allPersent/58) < 0.55:
        NAME_DOC_INPUT = 'LOWSertificat.docx'
        furst = 1
        second = 2
    else:
        NAME_DOC_INPUT = 'HightSertificat.docx'
        furst = 2
        second = 1

    doc = docx.Document(f'/Users/igorgerasimov/Python/DDKGTA/SertificatQUIZ/{NAME_DOC_INPUT}')
    #Глобальная настройка стиля  
    
    style = doc.styles['Normal']
    font = style.font
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)#font.name = 'Times New Roman'
    font.name = 'Jura'
    font.size = Pt(48)

    user3 =[]
    user3.append(split_List(user[7]))

    #Запись данных в Word
    tabelName = doc.tables[0]
    cellName = tabelName.cell(0,0)
    cellName.text = f'{user[0]}'

    #https://stackoverflow.com/questions/43007725/python-docx-how-to-change-table-font-size
    for row in tabelName.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(48)


    style = doc.styles['Normal']
    font = style.font
    font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) 
    font.name = 'Jura'
    font.size = Pt(20)
   
   # процент ответов
    tabelAnswer = doc.tables[furst]
    cellAnswer = tabelAnswer.cell(0,1)
    cellAnswer.text = f'{toFixed(int(user[1])/10 * 100)}%'

    cellAnswer = tabelAnswer.cell(1,1)
    cellAnswer.text = f'{toFixed(int(user[2])/10 * 100)}%'

    cellAnswer = tabelAnswer.cell(2,1)
    cellAnswer.text = f'{toFixed(int(user[3])/8 * 100)}%'

    cellAnswer = tabelAnswer.cell(3,1)
    cellAnswer.text = f'{toFixed(int(user[4])/8 * 100)}%'

    cellAnswer = tabelAnswer.cell(4,1)
    cellAnswer.text = f'{toFixed(int(user[5])/10 * 100)}%'

    cellAnswer = tabelAnswer.cell(5,1)
    cellAnswer.text = f'{toFixed(int(user[6])/10 * 100)}%'

# Cостав команды
    table = doc.tables[second]
    index = 0
    for item in user3[0]:
        cell = table.cell(index, 0)
        cell.text = f'{item}'
        index += 1

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(16) 
# Итоговый процент
    tablePersent = doc.tables[3]
    cellPersent = tablePersent.cell(0,0)
    allPersent = 0

    for index in range(1,7):
        allPersent += int(user[index])
    cellPersent.text = f'{toFixed((allPersent/56)*100)} %'
    
    
    for row in tablePersent.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(52) 

    doc.save(f'{user[int(NAME_DOC_OUT)-1]}.docx')

@logger.catch
def main():
    Clients = auntification_Sheet()
    sheet = Clients.open('QUIZ').worksheet('Лист1')
    for index in range(2,17):
        add_in_doc(sheet.get(f"A{index}:H{index}")[0]) # В ответ получаем как в TEST_LIST
        time.sleep(1)
   #add_in_doc(TEST_LIST[0], TEST_LIST2[0])
   # add_in_doc(sheet.get(f"A16:H16")[0]) # В ответ получаем как в TEST_LIST
 

if __name__ == "__main__":
    main()
   
